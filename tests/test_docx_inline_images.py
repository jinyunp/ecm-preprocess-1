import asyncio
import base64
from pathlib import Path

from docx import Document

from mypkg.components.parser.xml_parser import DocxXmlParser

PNG_BYTES = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Yprku4AAAAASUVORK5CYII="
)


def test_inline_image_placeholder_in_paragraph(tmp_path: Path) -> None:
    image_path = tmp_path / "inline.png"
    image_path.write_bytes(PNG_BYTES)

    doc = Document()
    paragraph = doc.add_paragraph("앞")
    paragraph.add_run().add_picture(str(image_path))
    paragraph.add_run("뒤")

    doc_path = tmp_path / "inline_image.docx"
    doc.save(doc_path)

    parser = DocxXmlParser()
    result = asyncio.run(parser.parse(doc_path))
    assert result.success

    payload = result.content
    paragraphs = payload["paragraphs"]
    inline_images = payload["inline_images"]

    assert inline_images, "인라인 이미지가 추출되지 않았습니다."
    rid = inline_images[0].rId
    placeholder = f"[image:{rid}]"

    assert paragraphs[0].text == f"앞{placeholder}뒤"
    assert placeholder in paragraphs[0].text
    assert inline_images[0].doc_index == paragraphs[0].doc_index


def test_paragraph_styles_from_xml(tmp_path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Heading", style="Heading 2")
    doc.add_paragraph("본문")

    doc_path = tmp_path / "styles.docx"
    doc.save(doc_path)

    parser = DocxXmlParser()
    result = asyncio.run(parser.parse(doc_path))
    assert result.success

    payload = result.content
    paragraphs = payload["paragraphs"]

    assert paragraphs[0].style.lower() == "heading 2"
    assert paragraphs[1].style is None
